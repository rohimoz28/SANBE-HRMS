from odoo import http
from odoo.http import request
from docx import Document
import io
import os

# Month Mapping
month_map = {
    '01': 'Januari', '02': 'Februari', '03': 'Maret',
    '04': 'April', '05': 'Mei', '06': 'Juni',
    '07': 'Juli', '08': 'Agustus', '09': 'September',
    '10': 'Oktober', '11': 'November', '12': 'Desember'
}

# Number Mapping
number_map = {
    '00': 'Nol', '01': 'Satu', '02': 'Dua', '03': 'Tiga', '04': 'Empat', '05': 'Lima',
    '06': 'Enam', '07': 'Tujuh', '08': 'Delapan', '09': 'Sembilan',
    '10': 'Sepuluh', '11': 'Sebelas', '12': 'Dua Belas', '13': 'Tiga Belas',
    '14': 'Empat Belas', '15': 'Lima Belas', '16': 'Enam Belas',
    '17': 'Tujuh Belas', '18': 'Delapan Belas', '19': 'Sembilan Belas',
    '20': 'Dua Puluh', '21': 'Dua Puluh Satu', '22': 'Dua Puluh Dua',
    '23': 'Dua Puluh Tiga', '24': 'Dua Puluh Empat', '25': 'Dua Puluh Lima',
    '26': 'Dua Puluh Enam', '27': 'Dua Puluh Tujuh', '28': 'Dua Puluh Delapan',
    '29': 'Dua Puluh Sembilan', '30': 'Tiga Puluh', '31': 'Tiga Puluh Satu'
}

# Year Helper
# Mengambil dari number map dari 1 - 19
year_helper_map = {
    int(k): v
    for k, v in number_map.items()
    if 1 <= int(k) <= 19
}

# Gender Mapping
gender_map = {
    'male': 'Laki-laki',
    'female': 'Perempuan',
    'other': '',
}

day_map = {
    'Monday': 'Senin',
    'Tuesday': 'Selasa',
    'Wednesday': 'Rabu',
    'Thursday': 'Kamis',
    'Friday': 'Jumat',
    'Saturday': 'Sabtu',
    'Sunday': 'Minggu',
}

class PKWTController(http.Controller):

    def spell_two_digits(self, n):
        if n < 20:
            return year_helper_map[n]
        tens = n // 10
        ones = n % 10
        text = year_helper_map[tens] + ' Puluh'
        if ones:
            text += ' ' + year_helper_map[ones]
        return text
    
    def spell_year_id(self, year: int) -> str:
        if year < 1000:
            return str(year)

        thousands = year // 1000
        remainder = year % 1000

        result = ''

        # ribu
        if thousands == 1:
            result = 'Seribu'
        else:
            result = year_helper_map[thousands] + ' Ribu'

        # ratus
        hundreds = remainder // 100
        tens = remainder % 100

        if hundreds:
            if hundreds == 1:
                result += ' Seratus'
            else:
                result += ' ' + year_helper_map[hundreds] + ' Ratus'

        if tens:
            result += ' ' + self.spell_two_digits(tens)

        return result.strip()

    def _format_date(self, record):
        if not record:
            return '-'
        day = record.strftime('%d')
        month = month_map[record.strftime('%m')]
        year = record.strftime('%Y')
        return f"{day} {month} {year}"

    def _replace_in_runs(self, paragraph, mapping):
        """
        Replace placeholder meskipun terbelah di beberapa run
        TANPA merusak font, size, bold, alignment, spasi
        """

        if not paragraph.runs:
            return

        for key, value in mapping.items():
            runs = paragraph.runs

            # Gabungkan teks + simpan index run
            full_text = ''
            run_map = []
            for i, run in enumerate(runs):
                run_map.append((i, len(full_text), len(full_text) + len(run.text)))
                full_text += run.text

            if key not in full_text:
                continue

            start = full_text.index(key)
            end = start + len(key)

            # Replace hanya run yang terkena placeholder
            for i, s, e in run_map:
                run = runs[i]

                # Run sepenuhnya di luar placeholder
                if e <= start or s >= end:
                    continue

                # Run terkena placeholder
                text = run.text
                new_text = ''

                for idx, char in enumerate(text):
                    global_pos = s + idx
                    if not (start <= global_pos < end):
                        new_text += char

                run.text = new_text

            # Tambahkan hasil ke run pertama placeholder (style ikut template)
            for i, s, e in run_map:
                if s <= start < e:
                    runs[i].text += value
                    break


    @http.route('/pkwt/download/<int:rec_id>', type='http', auth='user')
    def download_pkwt(self, rec_id):
        contract = request.env['hr.contract'].browse(rec_id)
        if not contract.exists():
            return request.not_found()

        # Path template
        base_dir = os.path.dirname(os.path.abspath(__file__))
        template_path = os.path.join(
            base_dir,
            '..',
            'static',
            'docx',
            'pkwt_template.docx'
        )

        # Load template
        doc = Document(template_path)

        month_diff = (contract.date_end.year - contract.date_start.year) * 12 + \
                     (contract.date_end.month - contract.date_start.month)
        
        month_diff_str = f"{month_diff:02d}"

        address = request.env['hr.employee.addresses'].search([
            ('employee_id','=', contract.employee_id.id),
            ('address_type','=', 'KTP'),
            ('default','=', True)
            ], limit=1)

        # Mapping placeholder
        mapping = {
            '{{ employee_name }}': contract.employee_id.name or '',
            '{{ employee_ktp }}': contract.employee_id.no_ktp or '',
            '{{ employee_gender }}': gender_map.get(contract.employee_id.gender, ''),
            '{{ employee_job }}': contract.job_id.with_context(lang='en_US').name or '',
            '{{ employee_birth }}': f"{contract.employee_id.place_of_birth or ''}, {self._format_date(contract.employee_id.birthday) or ''}",
            '{{ date_start }}': self._format_date(contract.date_start) or '',
            '{{ date_start_header }}': f"{contract.date_start.strftime('%d')}-{contract.date_start.strftime('%m')}-{contract.date_start.strftime('%Y')}" or '',
            '{{ date_end }}': self._format_date(contract.date_end) or '',
            '{{ month_date }}': month_map[contract.date_start.strftime('%m')] or '',
            '{{ date_spelled }}': number_map[contract.date_start.strftime('%d')] or '',
            '{{ month_diff_str }}': f"{month_diff_str} ({number_map[month_diff_str]})" or '',
            '{{ year_spelled }}': self.spell_year_id(int(contract.date_start.strftime('%Y'))) or '',
            '{{ address }}': f"{address.address or ''} Rt.{address.rt or ''} Rw.{address.rt or ''}" or '',
            '{{ city }}': f"Kel.{address.subdistrict or ''} Kec.{address.district or ''} {address.city or ''}" or '',
            '{{ day }}': day_map.get(contract.date_start.strftime('%A'), contract.date_start.strftime('%A')) or '',
            '{{ no_contract }}': contract.name or '',

        }

        # Replace di paragraf biasa
        for p in doc.paragraphs:
            self._replace_in_runs(p, mapping)

        # Replace di tabel
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        self._replace_in_runs(p, mapping)

        # Simpan ke memory
        buffer = io.BytesIO()
        doc.save(buffer)
        buffer.seek(0)

        filename = f'PKWT_{contract.employee_id.name}.docx'

        return request.make_response(
            buffer.read(),
            headers=[
                (
                    'Content-Type',
                    'application/vnd.openxmlformats-officedocument.wordprocessingml.document'
                ),
                (
                    'Content-Disposition',
                    f'attachment; filename="{filename}"'
                ),
            ]
        )
